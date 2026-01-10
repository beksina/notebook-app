"""Document processing service for text extraction from various formats."""

from pathlib import Path
import io

from PyPDF2 import PdfReader
from docx import Document


class DocumentProcessor:
    """Extract text from various document formats."""

    SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx"}

    def extract_text(self, file_path: Path) -> str:
        """Extract text content from document based on file type."""
        suffix = file_path.suffix.lower()

        if suffix == ".pdf":
            return self._extract_pdf(file_path)
        elif suffix in {".txt", ".md"}:
            return self._extract_text_file(file_path)
        elif suffix == ".docx":
            return self._extract_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

    def extract_text_from_bytes(self, content: bytes, filename: str) -> str:
        """Extract text from bytes based on filename extension."""
        suffix = Path(filename).suffix.lower()

        if suffix == ".pdf":
            return self._extract_pdf_from_bytes(content)
        elif suffix in {".txt", ".md"}:
            return content.decode("utf-8")
        elif suffix == ".docx":
            return self._extract_docx_from_bytes(content)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

    def _extract_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        reader = PdfReader(str(file_path))
        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        return "\n\n".join(text_parts)

    def _extract_pdf_from_bytes(self, content: bytes) -> str:
        """Extract text from PDF bytes."""
        reader = PdfReader(io.BytesIO(content))
        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        return "\n\n".join(text_parts)

    def _extract_text_file(self, file_path: Path) -> str:
        """Read text/markdown file."""
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

    def _extract_docx(self, file_path: Path) -> str:
        """Extract text from Word document."""
        doc = Document(str(file_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)

    def _extract_docx_from_bytes(self, content: bytes) -> str:
        """Extract text from Word document bytes."""
        doc = Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)


# Singleton
document_processor = DocumentProcessor()
